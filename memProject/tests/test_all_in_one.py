# -*- coding: utf-8 -*-
"""
记忆生成与去重融合模块 — 一键全功能测试脚本。

覆盖范围:
  Phase 1 — 单元测试 (Mock, 无外部依赖): LLMClient, EmbeddingClient,
           MemoryExtractor, MemoryGenerator, DedupService (24 项)
  Phase 2 — Schema 验证与映射 (无外部依赖): Pydantic 校验, Pipeline→Results 映射 (8 项)
  Phase 3 — E2E 基础 (需 DeepSeek + SiliconFlow API): LLM/Embedding 连接,
           真实抽取, 真实生成, 去重决策矩阵 (7 项)
  Phase 4 — E2E 完整 (需 Docker PostgreSQL + Qdrant): 完整流水线, DB 去重,
           MemoryStore CRUD (6 项)
  Phase 5 — 降级与容错: Qdrant 不可用, 空抽取, 关键词提取边界 (5 项)

用法:
  # 仅 Phase 1+2（无外部依赖）
  python tests/test_all_in_one.py --mock-only

  # Phase 1-3（不依赖数据库）
  python tests/test_all_in_one.py --no-db

  # 全部测试（需 Docker + API）
  python tests/test_all_in_one.py

  # 全部测试 + 生成 Word 报告
  python tests/test_all_in_one.py --report

输出:
  - 控制台实时输出
  - tests/test_all_in_one_results.json（结构化结果）
  - tests/记忆生成与去重融合_全功能测试报告.docx（--report 时生成）
"""

import asyncio
import json
import os
import sys
import time
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from unittest.mock import AsyncMock, MagicMock, patch

# === 环境设置 ===
os.environ.setdefault(
    "SSL_CERT_FILE",
    "E:/anaconda3/envs/agent_mem/Lib/site-packages/certifi/cacert.pem",
)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_PROJECT_ROOT))
sys.stdout.reconfigure(encoding="utf-8")

# === 结果收集 ===
RESULTS: list[dict] = []
START_TIME = datetime.now()
TEST_USER_ID = f"test_all_in_one_{int(time.time())}"


def record(
    phase: str,
    name: str,
    status: str,  # "pass" | "fail" | "skip"
    duration_ms: float = 0,
    details: str = "",
    error: str = "",
) -> None:
    RESULTS.append({
        "phase": phase,
        "name": name,
        "status": status,
        "duration_ms": round(duration_ms, 1),
        "details": details,
        "error": error,
        "timestamp": datetime.now().isoformat(),
    })


def divider(title: str = "") -> None:
    width = 70
    if title:
        print(f"\n{'=' * width}")
        print(f"  {title}")
        print(f"{'=' * width}")
    else:
        print(f"{'─' * width}")


# ====================================================================
# Phase 1: 单元测试 (Mock 外部依赖)
# ====================================================================

async def phase1_unit_tests():
    """运行全部 24 个单元测试，使用 Mock 替代外部依赖。"""
    divider("Phase 1: 单元测试 (Mock 模式)")

    # ---- 1.1 LLM Client ----
    from app.services.llm_client import LLMClient

    print("\n--- 1.1 LLMClient (4 tests) ---")

    llm = LLMClient(api_key="test-key", base_url="https://test.api.com/v1", model="test-model")

    # Test 1: chat_completion success
    t0 = time.perf_counter()
    try:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.json.return_value = {"choices": [{"message": {"content": "Hello!"}}]}
        mock_resp.raise_for_status = MagicMock()
        mock_http = AsyncMock()
        mock_http.post = AsyncMock(return_value=mock_resp)
        llm._http = mock_http

        result = await llm.chat_completion(messages=[{"role": "user", "content": "Hi"}])
        assert result == "Hello!", f"Unexpected: {result}"
        record("Phase1-LLM", "chat_completion_success", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_chat_completion_success")
    except Exception as e:
        record("Phase1-LLM", "chat_completion_success", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_chat_completion_success: {e}")

    # Test 2: extract_structured
    t0 = time.perf_counter()
    try:
        llm.chat_completion = AsyncMock(return_value='{"name": "test", "value": 42}')
        result = await llm.extract_structured("prompt", "content", {"type": "object"})
        assert result == {"name": "test", "value": 42}
        record("Phase1-LLM", "extract_structured", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_extract_structured")
    except Exception as e:
        record("Phase1-LLM", "extract_structured", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_extract_structured: {e}")

    # Test 3: JSON recovery
    t0 = time.perf_counter()
    try:
        llm.chat_completion = AsyncMock(return_value='Prefix text {"name": "recovered"} suffix')
        result = await llm.extract_structured("prompt", "content", {"type": "object"})
        assert result == {"name": "recovered"}
        record("Phase1-LLM", "extract_structured_json_recovery", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_extract_structured_json_recovery")
    except Exception as e:
        record("Phase1-LLM", "extract_structured_json_recovery", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_extract_structured_json_recovery: {e}")

    # Test 4: Complete failure
    t0 = time.perf_counter()
    try:
        from app.core.exceptions import LLMServiceError
        llm.chat_completion = AsyncMock(return_value="No JSON here at all")
        try:
            await llm.extract_structured("prompt", "content", {"type": "object"})
            record("Phase1-LLM", "extract_structured_complete_failure", "fail", (time.perf_counter() - t0) * 1000,
                   error="Should have raised LLMServiceError")
            print("  ❌ test_extract_structured_complete_failure: no exception raised")
        except LLMServiceError:
            record("Phase1-LLM", "extract_structured_complete_failure", "pass", (time.perf_counter() - t0) * 1000)
            print("  ✅ test_extract_structured_complete_failure")
    except Exception as e:
        record("Phase1-LLM", "extract_structured_complete_failure", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_extract_structured_complete_failure: {e}")

    # ---- 1.2 Embedding Client ----
    from app.services.embedding_client import EmbeddingClient

    print("\n--- 1.2 EmbeddingClient (2 tests) ---")

    emb = EmbeddingClient(api_key="test-key", base_url="https://test.api.com/v1", model="test-model")

    # Test 5: embed_single
    t0 = time.perf_counter()
    try:
        mock_http = AsyncMock()
        mock_http.post = AsyncMock(return_value=MagicMock(
            status_code=200,
            json=MagicMock(return_value={"data": [{"embedding": [0.1] * 1024, "index": 0}]}),
            raise_for_status=MagicMock(),
        ))
        emb._http = mock_http
        result = await emb.embed_single("test text")
        assert len(result) == 1024
        assert result[0] == 0.1
        record("Phase1-Emb", "embed_single", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_embed_single")
    except Exception as e:
        record("Phase1-Emb", "embed_single", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_embed_single: {e}")

    # Test 6: embed_batch splitting
    t0 = time.perf_counter()
    try:
        call_count = 0
        async def mock_post(url, json=None):
            nonlocal call_count
            call_count += 1
            texts = json.get("input", [])
            return MagicMock(
                status_code=200,
                json=MagicMock(return_value={
                    "data": [{"embedding": [0.5] * 1024, "index": i} for i in range(len(texts))],
                }),
                raise_for_status=MagicMock(),
            )
        mock_http = AsyncMock()
        mock_http.post = mock_post
        emb._http = mock_http
        results = await emb.embed_batch(["text"] * 40)
        assert len(results) == 40
        assert call_count == 2
        record("Phase1-Emb", "embed_batch_splitting", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_embed_batch_splitting")
    except Exception as e:
        record("Phase1-Emb", "embed_batch_splitting", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_embed_batch_splitting: {e}")

    # ---- 1.3 MemoryExtractor ----
    from app.services.memory_extractor import (
        MemoryExtractor, KeyFactExtractor, TaskStateExtractor, DecisionExtractor,
        ExtractionResult, KeyFactsResult, TaskStateResult, DecisionResult,
    )

    print("\n--- 1.3 MemoryExtractor (5 tests) ---")

    mock_llm = MagicMock(spec=LLMClient)

    # Test 7: KeyFact extraction
    t0 = time.perf_counter()
    try:
        mock_llm.extract_structured = AsyncMock(return_value={
            "business_objects": [{"name": "ProjectX", "type": "project", "description": "main"}],
            "constraints": [{"type": "temporal", "description": "deadline Friday", "severity": "high"}],
            "confirmations": [{"item": "Use FastAPI", "parties": ["dev"], "context": "meeting"}],
        })
        extractor = KeyFactExtractor(mock_llm)
        result = await extractor.extract("test text")
        assert len(result.business_objects) == 1
        assert result.business_objects[0]["name"] == "ProjectX"
        assert len(result.constraints) == 1
        assert len(result.confirmations) == 1
        record("Phase1-Ext", "key_fact_extraction", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_key_fact_extraction")
    except Exception as e:
        record("Phase1-Ext", "key_fact_extraction", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_key_fact_extraction: {e}")

    # Test 8: TaskState extraction
    t0 = time.perf_counter()
    try:
        mock_llm.extract_structured = AsyncMock(return_value={
            "current_progress": "API development phase",
            "completed_items": [{"item": "DB design", "evidence": "ER diagram done"}],
            "pending_items": [{"item": "Unit tests", "priority": "high"}],
        })
        extractor = TaskStateExtractor(mock_llm)
        result = await extractor.extract("test text")
        assert "API development" in result.current_progress
        assert len(result.completed_items) == 1
        assert len(result.pending_items) == 1
        record("Phase1-Ext", "task_state_extraction", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_task_state_extraction")
    except Exception as e:
        record("Phase1-Ext", "task_state_extraction", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_task_state_extraction: {e}")

    # Test 9: Decision extraction
    t0 = time.perf_counter()
    try:
        mock_llm.extract_structured = AsyncMock(return_value={
            "confirmed_plans": [{"plan": "Microservices", "alternatives": ["Monolith"], "decision_context": "review"}],
            "selection_rationale": [{"reason": "Better scalability", "criteria": ["scalability", "maintainability"]}],
            "execution_results": [{"result": "Deploy success", "outcome_type": "success"}],
        })
        extractor = DecisionExtractor(mock_llm)
        result = await extractor.extract("test text")
        assert len(result.confirmed_plans) == 1
        assert result.confirmed_plans[0]["plan"] == "Microservices"
        assert len(result.selection_rationale) == 1
        assert len(result.execution_results) == 1
        record("Phase1-Ext", "decision_extraction", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_decision_extraction")
    except Exception as e:
        record("Phase1-Ext", "decision_extraction", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_decision_extraction: {e}")

    # Test 10: MemoryExtractor facade
    t0 = time.perf_counter()
    try:
        mock_llm.extract_structured = AsyncMock(return_value={
            "business_objects": [], "constraints": [], "confirmations": [],
        })
        extractor = MemoryExtractor(mock_llm)
        result = await extractor.extract(text="test", types=["key_fact"])
        assert result.key_facts is not None
        assert result.task_state is None
        assert result.decisions is None
        record("Phase1-Ext", "memory_extractor_facade", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_memory_extractor_facade")
    except Exception as e:
        record("Phase1-Ext", "memory_extractor_facade", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_memory_extractor_facade: {e}")

    # Test 11: ExtractionResult.to_dict
    t0 = time.perf_counter()
    try:
        result = ExtractionResult(
            key_facts=KeyFactsResult(
                business_objects=[{"name": "Test"}],
                constraints=[],
                confirmations=[],
            ),
            source_text="test",
        )
        d = result.to_dict()
        assert len(d["business_objects"]) == 1
        assert d["constraints"] == []
        assert d["current_progress"] == ""
        record("Phase1-Ext", "extraction_result_to_dict", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_extraction_result_to_dict")
    except Exception as e:
        record("Phase1-Ext", "extraction_result_to_dict", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_extraction_result_to_dict: {e}")

    # ---- 1.4 MemoryGenerator ----
    from app.services.memory_generator import MemoryGenerator, MemoryCandidate

    print("\n--- 1.4 MemoryGenerator (3 tests) ---")

    # Test 12: generate_memories
    t0 = time.perf_counter()
    try:
        mock_llm.extract_structured = AsyncMock(return_value={
            "memories": [
                {
                    "content": "User prefers Python",
                    "summary": "Python preference",
                    "key_points": ["Python", "FastAPI"],
                    "memory_type": "preference",
                    "tags": ["python"],
                    "entities": ["Python"],
                    "importance": 0.8,
                    "confidence": 0.9,
                },
                {
                    "content": "ProjectX deadline is Friday",
                    "summary": "Deadline",
                    "key_points": ["deadline: Friday"],
                    "memory_type": "constraint",
                    "tags": ["deadline"],
                    "entities": ["ProjectX"],
                    "importance": 0.9,
                    "confidence": 1.0,
                },
            ]
        })
        extraction = ExtractionResult(
            key_facts=KeyFactsResult(
                business_objects=[{"name": "ProjectX", "type": "project", "description": ""}],
                constraints=[],
                confirmations=[{"item": "Use Python", "parties": ["user"], "context": ""}],
            ),
            source_text="test",
        )
        generator = MemoryGenerator(mock_llm)
        candidates = await generator.generate(extraction)
        assert len(candidates) == 2
        assert candidates[0].memory_type == "preference"
        assert candidates[1].memory_type == "constraint"
        record("Phase1-Gen", "generate_memories", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_generate_memories")
    except Exception as e:
        record("Phase1-Gen", "generate_memories", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_generate_memories: {e}")

    # Test 13: generate_empty_extraction
    t0 = time.perf_counter()
    try:
        extraction = ExtractionResult(source_text="hello")
        generator = MemoryGenerator(mock_llm)
        candidates = await generator.generate(extraction)
        assert len(candidates) == 0
        record("Phase1-Gen", "generate_empty_extraction", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_generate_empty_extraction")
    except Exception as e:
        record("Phase1-Gen", "generate_empty_extraction", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_generate_empty_extraction: {e}")

    # Test 14: memory_candidate_validation
    t0 = time.perf_counter()
    try:
        c = MemoryCandidate(content="test", summary="s", memory_type="fact", importance=0.7, confidence=0.8)
        c.validate()
        assert c.memory_type == "fact"

        c2 = MemoryCandidate(content="test", summary="s", memory_type="invalid_type")
        c2.validate()
        assert c2.memory_type == "fact"

        c3 = MemoryCandidate(content="test", summary="s", importance=1.5, confidence=-0.5)
        c3.validate()
        assert c3.importance == 1.0
        assert c3.confidence == 0.0
        record("Phase1-Gen", "memory_candidate_validation", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_memory_candidate_validation")
    except Exception as e:
        record("Phase1-Gen", "memory_candidate_validation", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_memory_candidate_validation: {e}")

    # ---- 1.5 DedupService ----
    from app.services.memory_dedup import DedupService, DedupAction, DedupResult, SimilarMemory
    from app.models.base import Memory

    print("\n--- 1.5 DedupService (10 tests) ---")

    mock_embedding = MagicMock(spec=EmbeddingClient)
    mock_embedding.embed_single = AsyncMock(return_value=[0.1] * 1024)
    mock_qdrant = MagicMock()
    mock_qdrant.is_available = True
    mock_qdrant.collection_name = "agent_mem_generation"

    # Test 15: decide_action_discard
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        action = service._decide_action(0.95, 0.85, True)
        assert action == DedupAction.DISCARD
        record("Phase1-Dedup", "decide_action_discard", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_decide_action_discard")
    except Exception as e:
        record("Phase1-Dedup", "decide_action_discard", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_decide_action_discard: {e}")

    # Test 16: decide_action_update_existing
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        action = service._decide_action(0.90, 0.70, True)
        assert action == DedupAction.UPDATE_EXISTING
        record("Phase1-Dedup", "decide_action_update_existing", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_decide_action_update_existing")
    except Exception as e:
        record("Phase1-Dedup", "decide_action_update_existing", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_decide_action_update_existing: {e}")

    # Test 17: decide_action_merge
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        action = service._decide_action(0.88, 0.76, False)
        assert action == DedupAction.MERGE
        record("Phase1-Dedup", "decide_action_merge", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_decide_action_merge")
    except Exception as e:
        record("Phase1-Dedup", "decide_action_merge", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_decide_action_merge: {e}")

    # Test 18: decide_action_keep_new
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        action = service._decide_action(0.60, 0.20, False)
        assert action == DedupAction.KEEP_NEW
        record("Phase1-Dedup", "decide_action_keep_new", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_decide_action_keep_new")
    except Exception as e:
        record("Phase1-Dedup", "decide_action_keep_new", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_decide_action_keep_new: {e}")

    # Test 19: compute_keyword_overlap
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        candidate = MemoryCandidate(
            content="Use Python and FastAPI",
            tags=["python", "fastapi"], entities=["Python", "FastAPI"],
        )
        existing = Memory(
            content="Python project with Django",
            tags=["python", "django"], entities=["Python", "Django"],
        )
        overlap = service._compute_keyword_overlap(candidate, existing)
        assert 0.0 <= overlap <= 1.0
        record("Phase1-Dedup", "compute_keyword_overlap", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_compute_keyword_overlap")
    except Exception as e:
        record("Phase1-Dedup", "compute_keyword_overlap", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_compute_keyword_overlap: {e}")

    # Test 20: check_identity_same_task
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        candidate = MemoryCandidate(content="test", entities=[])
        existing = Memory(task_id="task_001", content="test", entities=[])
        assert service._check_identity(candidate, existing, task_id="task_001")
        assert not service._check_identity(candidate, existing, task_id="task_002")
        record("Phase1-Dedup", "check_identity_same_task", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_check_identity_same_task")
    except Exception as e:
        record("Phase1-Dedup", "check_identity_same_task", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_check_identity_same_task: {e}")

    # Test 21: check_identity_entity_overlap
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        candidate = MemoryCandidate(content="test", entities=["Python", "FastAPI", "Qdrant"])
        existing = Memory(content="test", entities=["Python", "FastAPI", "Django"], task_id=None)
        assert service._check_identity(candidate, existing)
        assert not service._check_identity(
            MemoryCandidate(content="t", entities=["X"]),
            Memory(content="t", entities=["Y"]),
        )
        record("Phase1-Dedup", "check_identity_entity_overlap", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_check_identity_entity_overlap")
    except Exception as e:
        record("Phase1-Dedup", "check_identity_entity_overlap", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_check_identity_entity_overlap: {e}")

    # Test 22: merge_content
    t0 = time.perf_counter()
    try:
        service = DedupService(mock_embedding, mock_qdrant)
        candidate = MemoryCandidate(
            content="New: use FastAPI", summary="FastAPI choice",
            key_points=["FastAPI"], tags=["fastapi"], entities=["FastAPI"],
            importance=0.9, confidence=0.95,
        )
        existing = Memory(
            content="User uses Python", summary="Python pref",
            key_points=["Python"], tags=["python"], entities=["Python"],
            importance=0.7, confidence=0.8,
        )
        merged = service._merge_content(candidate, existing)
        assert "Python" in merged["content"]
        assert "FastAPI" in merged["content"]
        assert "fastapi" in merged["entities"]
        assert "python" in merged["entities"]
        assert merged["importance"] == 0.9
        assert merged["confidence"] == 0.95
        record("Phase1-Dedup", "merge_content", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_merge_content")
    except Exception as e:
        record("Phase1-Dedup", "merge_content", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_merge_content: {e}")

    # Test 23: process_candidates_qdrant_unavailable
    t0 = time.perf_counter()
    try:
        mock_qdrant.is_available = False
        service = DedupService(mock_embedding, mock_qdrant)
        candidates = [
            MemoryCandidate(content="test1", summary="s1"),
            MemoryCandidate(content="test2", summary="s2"),
        ]
        db = AsyncMock()
        results = await service.process_candidates(candidates, "user_001", db)
        assert len(results) == 2
        assert all(r.action == DedupAction.KEEP_NEW for r in results)
        mock_qdrant.is_available = True
        record("Phase1-Dedup", "process_candidates_qdrant_unavailable", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_process_candidates_qdrant_unavailable")
    except Exception as e:
        mock_qdrant.is_available = True
        record("Phase1-Dedup", "process_candidates_qdrant_unavailable", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_process_candidates_qdrant_unavailable: {e}")

    # Test 24: extract_nouns
    t0 = time.perf_counter()
    try:
        text = "用户使用 Python 开发 Web 应用，偏好 FastAPI 框架"
        nouns = DedupService._extract_nouns(text)
        has_python = any("python" in n.lower() or "Python" in n for n in nouns)
        assert has_python or len(nouns) > 0, "Should extract some keywords"
        record("Phase1-Dedup", "extract_nouns", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ test_extract_nouns")
    except Exception as e:
        record("Phase1-Dedup", "extract_nouns", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ test_extract_nouns: {e}")


# ====================================================================
# Phase 2: Schema 验证与映射
# ====================================================================

async def phase2_schema_and_mapping():
    """测试 Pydantic Schema 验证和 Pipeline → 前端 Results 映射。"""
    divider("Phase 2: Schema 验证与映射")

    # ---- 2.1 Schema Validation ----
    print("\n--- 2.1 Schema Validation (7 tests) ---")

    from app.schemas.memory import MemoryWriteRequest, MessageItem

    # Test 25: Normal request
    t0 = time.perf_counter()
    try:
        body = MemoryWriteRequest(
            user_id=TEST_USER_ID,
            scene_id="chat",
            messages=[
                MessageItem(role="user", content="我叫张伟，喜欢Python"),
                MessageItem(role="assistant", content="你好张伟！"),
            ],
        )
        assert body.user_id == TEST_USER_ID
        assert len(body.messages) == 2
        assert "张伟" in body.get_content_text()
        record("Phase2-Schema", "normal_request_validation", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ normal_request_validation")
    except Exception as e:
        record("Phase2-Schema", "normal_request_validation", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ normal_request_validation: {e}")

    # Test 26: Missing user_id
    t0 = time.perf_counter()
    try:
        MemoryWriteRequest(messages=[MessageItem(role="user", content="test")])
        record("Phase2-Schema", "missing_user_id_rejected", "fail", (time.perf_counter() - t0) * 1000,
               error="Should have raised ValidationError")
        print("  ❌ missing_user_id_rejected: no exception")
    except Exception:
        record("Phase2-Schema", "missing_user_id_rejected", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ missing_user_id_rejected")

    # Test 27: Invalid role
    t0 = time.perf_counter()
    try:
        MessageItem(role="invalid_role", content="test")
        record("Phase2-Schema", "invalid_role_rejected", "fail", (time.perf_counter() - t0) * 1000,
               error="Should have raised ValidationError")
        print("  ❌ invalid_role_rejected: no exception")
    except Exception:
        record("Phase2-Schema", "invalid_role_rejected", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ invalid_role_rejected")

    # Test 28: Empty content
    t0 = time.perf_counter()
    try:
        MessageItem(role="user", content="")
        record("Phase2-Schema", "empty_content_rejected", "fail", (time.perf_counter() - t0) * 1000,
               error="Should have raised ValidationError")
        print("  ❌ empty_content_rejected: no exception")
    except Exception:
        record("Phase2-Schema", "empty_content_rejected", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ empty_content_rejected")

    # Test 29: Empty messages array
    t0 = time.perf_counter()
    try:
        MemoryWriteRequest(user_id="test", messages=[])
        record("Phase2-Schema", "empty_messages_rejected", "fail", (time.perf_counter() - t0) * 1000,
               error="Should have raised ValidationError")
        print("  ❌ empty_messages_rejected: no exception")
    except Exception:
        record("Phase2-Schema", "empty_messages_rejected", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ empty_messages_rejected")

    # Test 30: user_id auto normalize
    t0 = time.perf_counter()
    try:
        body = MemoryWriteRequest(
            user_id="  Test_User_02  ",
            messages=[MessageItem(role="user", content="test")],
        )
        assert body.user_id == "test_user_02", f"Expected 'test_user_02', got '{body.user_id}'"
        record("Phase2-Schema", "user_id_auto_normalize", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ user_id_auto_normalize")
    except Exception as e:
        record("Phase2-Schema", "user_id_auto_normalize", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ user_id_auto_normalize: {e}")

    # Test 31: GenerationRequest invalid extraction_type
    t0 = time.perf_counter()
    try:
        from app.schemas.generation import GenerationRequest
        GenerationRequest(text="test", user_id="u1", extraction_types=["invalid_type"])
        record("Phase2-Schema", "invalid_extraction_type_rejected", "fail", (time.perf_counter() - t0) * 1000,
               error="Should have raised ValidationError")
        print("  ❌ invalid_extraction_type_rejected: no exception")
    except Exception:
        record("Phase2-Schema", "invalid_extraction_type_rejected", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ invalid_extraction_type_rejected")

    # ---- 2.2 Pipeline → Results Mapping ----
    print("\n--- 2.2 Pipeline → Results Mapping (1 test) ---")

    # Test 32: Pipeline result mapping
    t0 = time.perf_counter()
    try:
        from app.api.v1.memory import _pipeline_to_write_results
        from app.services.memory_pipeline import PipelineResult

        result = PipelineResult(
            memory_ids=["mem_001", "mem_002", "mem_003", "mem_004"],
            new_count=2, merged_count=1, discarded_count=1, updated_count=1,
            details=[
                {"action": "keep_new", "memory_id": "mem_001", "content_preview": "A", "memory_type": "preference", "importance": 0.9, "confidence": 1.0, "message": "new"},
                {"action": "merge", "memory_id": "mem_002", "content_preview": "B", "memory_type": "fact", "importance": 0.8, "confidence": 1.0, "message": "merged"},
                {"action": "discard", "memory_id": "", "content_preview": "C", "memory_type": "fact", "importance": 0.1, "confidence": 0.3, "message": "skip"},
                {"action": "update_existing", "memory_id": "mem_003", "content_preview": "D", "memory_type": "constraint", "importance": 0.8, "confidence": 0.9, "message": "updated"},
            ],
        )

        mapped = _pipeline_to_write_results(result)
        assert len(mapped) == 4
        assert mapped[0].event.value == "ADD"
        assert mapped[0].id == "mem_001"
        assert mapped[1].event.value == "MERGE"
        assert mapped[2].event.value == "SKIP"
        assert mapped[2].id == ""
        assert mapped[3].event.value == "ADD"
        record("Phase2-Map", "pipeline_to_results_mapping", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ pipeline_to_results_mapping")
    except Exception as e:
        record("Phase2-Map", "pipeline_to_results_mapping", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ pipeline_to_results_mapping: {e}")


# ====================================================================
# Phase 3: E2E 基础 (真实 API)
# ====================================================================

SAMPLE_TEXT = """
用户: 我们正在开发一个叫 ProjectX 的电商平台，主要用 Python 和 FastAPI。
智能体: 好的，Python + FastAPI 技术栈已经记录。
用户: 项目的 deadline 是下周五，目前数据库设计已经完成，API 接口开发正在进行中。
智能体: 了解。数据库设计已完成，API 开发进行中。
用户: 我们之前讨论了技术选型，最终决定用 PostgreSQL 而不是 MySQL，因为需要 pgvector 支持。
智能体: 确认，已选择 PostgreSQL 作为数据库。
用户: 用户画像服务需要在下周三之前完成，QA 团队已经确认了这个时间点。
智能体: 记录：用户画像服务 deadline 为下周三，QA 已确认。
"""


async def phase3_e2e_basic():
    """E2E 基础测试 — 需要 DeepSeek + SiliconFlow API 连通。"""
    divider("Phase 3: E2E 基础测试 (真实 API)")

    from app.services.llm_client import llm_client
    from app.services.embedding_client import embedding_client
    from app.services.memory_extractor import MemoryExtractor
    from app.services.memory_generator import MemoryGenerator
    from app.services.memory_dedup import DedupService, DedupAction
    from app.core.qdrant_client import qdrant_client

    # ---- 3.1 LLM Connection ----
    print("\n--- 3.1 LLM 连接测试 ---")
    t0 = time.perf_counter()
    try:
        response = await llm_client.chat_completion(
            messages=[{"role": "user", "content": "Say 'OK' if you can read this."}],
            max_tokens=20,
        )
        ok = "OK" in response.upper()
        record("Phase3-E2E", "llm_connection", "pass" if ok else "fail",
               (time.perf_counter() - t0) * 1000,
               details=f"Response: {response.strip()[:100]}")
        print(f"  {'✅' if ok else '❌'} llm_connection: {response.strip()[:80]}")
        if not ok:
            print("  ⚠️ LLM API 不可用，跳过 Phase 3 后续测试")
            return False
    except Exception as e:
        record("Phase3-E2E", "llm_connection", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ llm_connection: {e}")
        print("  ⚠️ 跳过 Phase 3 后续测试")
        return False

    # ---- 3.2 Embedding Connection ----
    print("\n--- 3.2 Embedding 连接测试 ---")
    t0 = time.perf_counter()
    try:
        vec = await embedding_client.embed_single("Hello, world!")
        dim_ok = len(vec) == 1024
        record("Phase3-E2E", "embedding_connection", "pass" if dim_ok else "fail",
               (time.perf_counter() - t0) * 1000,
               details=f"Dimension: {len(vec)}, first 5: {[round(v, 4) for v in vec[:5]]}")
        print(f"  {'✅' if dim_ok else '❌'} embedding_connection: dim={len(vec)}")
    except Exception as e:
        record("Phase3-E2E", "embedding_connection", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ embedding_connection: {e}")

    # ---- 3.3 Real Extraction ----
    print("\n--- 3.3 真实 LLM 抽取 ---")
    extractor = MemoryExtractor(llm_client)

    t0 = time.perf_counter()
    try:
        result = await extractor.extract(SAMPLE_TEXT, types=["key_fact", "task_state", "decision"])
        elapsed = time.perf_counter() - t0
        kf_items = len(result.key_facts.business_objects) + len(result.key_facts.constraints) + len(result.key_facts.confirmations) if result.key_facts else 0
        ts_items = (len(result.task_state.completed_items) + len(result.task_state.pending_items)) if result.task_state else 0
        dc_items = len(result.decisions.confirmed_plans) + len(result.decisions.selection_rationale) + len(result.decisions.execution_results) if result.decisions else 0

        details = (f"KeyFacts: {kf_items} items, TaskState: {ts_items} items, "
                   f"Decisions: {dc_items} items, elapsed: {elapsed:.1f}s")
        has_content = kf_items + ts_items + dc_items > 0
        record("Phase3-E2E", "real_extraction", "pass" if has_content else "fail",
               elapsed * 1000, details=details)
        print(f"  {'✅' if has_content else '❌'} real_extraction ({elapsed:.1f}s): {kf_items}+{ts_items}+{dc_items} items")
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase3-E2E", "real_extraction", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ real_extraction: {e}")

    # ---- 3.4 Real Generation ----
    print("\n--- 3.4 真实 LLM 记忆生成 ---")
    t0 = time.perf_counter()
    try:
        extractor2 = MemoryExtractor(llm_client)
        extraction = await extractor2.extract(SAMPLE_TEXT)
        generator = MemoryGenerator(llm_client)
        candidates = await generator.generate(extraction)
        elapsed = time.perf_counter() - t0
        types_found = list(set(c.memory_type for c in candidates))
        details = (f"Generated {len(candidates)} candidates, types: {types_found}, "
                   f"elapsed: {elapsed:.1f}s")
        ok = len(candidates) >= 1
        record("Phase3-E2E", "real_generation", "pass" if ok else "fail",
               elapsed * 1000, details=details)
        print(f"  {'✅' if ok else '❌'} real_generation ({elapsed:.1f}s): {len(candidates)} candidates")
        for c in candidates[:3]:
            print(f"     - [{c.memory_type}] {c.content[:80]}...")
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase3-E2E", "real_generation", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ real_generation: {e}")

    # ---- 3.5 Dedup Decision Matrix ----
    print("\n--- 3.5 去重决策矩阵验证 ---")
    t0 = time.perf_counter()
    try:
        service = DedupService(embedding_client, qdrant_client)
        test_cases = [
            (0.95, 0.90, True, DedupAction.DISCARD),
            (0.90, 0.70, True, DedupAction.UPDATE_EXISTING),
            (0.88, 0.76, False, DedupAction.MERGE),
            (0.60, 0.30, False, DedupAction.KEEP_NEW),
        ]
        all_pass = True
        for vec, kw, identity, expected in test_cases:
            action = service._decide_action(vec, kw, identity)
            if action != expected:
                all_pass = False
            print(f"     vec={vec:.2f} kw={kw:.2f} id={identity} → {action.value} {'✅' if action == expected else '❌'}")
        record("Phase3-E2E", "dedup_decision_matrix", "pass" if all_pass else "fail",
               (time.perf_counter() - t0) * 1000,
               details=f"4/4 correct" if all_pass else "Some mismatches")
        print(f"  {'✅' if all_pass else '❌'} dedup_decision_matrix")
    except Exception as e:
        record("Phase3-E2E", "dedup_decision_matrix", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ dedup_decision_matrix: {e}")

    # ---- 3.6 Composite Score Edge Cases ----
    print("\n--- 3.6 复合评分边界测试 ---")
    t0 = time.perf_counter()
    try:
        service = DedupService(embedding_client, qdrant_client)
        edge_cases = [
            (0.85, 0.55, False, DedupAction.KEEP_NEW, "边界: 接近但不足 0.65"),
            (1.00, 1.00, True, DedupAction.DISCARD, "完全匹配"),
            (0.00, 0.00, False, DedupAction.KEEP_NEW, "完全无关"),
            (0.70, 0.50, True, DedupAction.MERGE, "边界: composite=0.70, identity → MERGE"),
        ]
        all_edge_pass = True
        for vec, kw, identity, expected, desc in edge_cases:
            action = service._decide_action(vec, kw, identity)
            if action != expected:
                all_edge_pass = False
            print(f"     {desc}: vec={vec:.2f} kw={kw:.2f} id={identity} → {action.value} {'✅' if action == expected else '❌'}")
        record("Phase3-E2E", "composite_score_edge_cases", "pass" if all_edge_pass else "fail",
               (time.perf_counter() - t0) * 1000,
               details=f"4/4 correct" if all_edge_pass else "Some mismatches")
        print(f"  {'✅' if all_edge_pass else '❌'} composite_score_edge_cases")
    except Exception as e:
        record("Phase3-E2E", "composite_score_edge_cases", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ composite_score_edge_cases: {e}")

    # Close clients
    await llm_client.close()
    await embedding_client.close()
    return True


# ====================================================================
# Phase 4: E2E 完整 (数据库 + Qdrant)
# ====================================================================

async def phase4_e2e_full():
    """E2E 完整测试 — 需要 Docker PostgreSQL + Qdrant。"""
    divider("Phase 4: E2E 完整测试 (数据库 + Qdrant)")

    from app.core.database import get_db, check_db_connection
    from app.core.qdrant_client import qdrant_client
    from app.services.memory_pipeline import memory_pipeline
    from app.services.memory_store import MemoryStore
    from app.models.base import Memory

    # Check DB
    db_ok = await check_db_connection()
    if not db_ok:
        print("  ⚠️ 数据库不可用，跳过 Phase 4 全部测试")
        print("     启动方式: docker start mem-postgres mem-qdrant")
        for name in ["full_pipeline", "dedup_with_db", "memory_store_search",
                      "memory_store_crud", "memory_store_context", "memory_store_update_delete"]:
            record("Phase4-DB", name, "skip", details="Database unavailable")
        return

    # Init Qdrant
    if not qdrant_client.is_available:
        print("  初始化 Qdrant...")
        ok = qdrant_client.initialize()
        if not ok:
            print("  ⚠️ Qdrant 初始化失败，跳过 Phase 4")
            for name in ["full_pipeline", "dedup_with_db", "memory_store_search",
                          "memory_store_crud", "memory_store_context", "memory_store_update_delete"]:
                record("Phase4-DB", name, "skip", details="Qdrant unavailable")
            return
        print("  ✅ Qdrant 连接成功")

    # ---- 4.1 Full Pipeline ----
    print("\n--- 4.1 完整流水线 (extract → generate → dedup → store) ---")
    t0 = time.perf_counter()
    try:
        async for db in get_db():
            result = await memory_pipeline.run(
                text=SAMPLE_TEXT,
                user_id=TEST_USER_ID,
                db=db,
            )
            elapsed = time.perf_counter() - t0
            details = (f"new={result.new_count}, merged={result.merged_count}, "
                       f"discarded={result.discarded_count}, updated={result.updated_count}, "
                       f"elapsed={elapsed:.1f}s")
            ok = result.new_count > 0
            record("Phase4-DB", "full_pipeline", "pass" if ok else "fail", elapsed * 1000, details=details)
            print(f"  {'✅' if ok else '❌'} full_pipeline ({elapsed:.1f}s): {details}")
            for d in result.details[:3]:
                print(f"     [{d['action']}] {d.get('content_preview', '')[:80]}")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "full_pipeline", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ full_pipeline: {e}")

    # ---- 4.2 Dedup with DB ----
    print("\n--- 4.2 去重测试 (相同内容两次提交) ---")
    t0 = time.perf_counter()
    try:
        dedup_user = f"test_dedup_{int(time.time())}"
        async for db in get_db():
            # First submission
            result1 = await memory_pipeline.run(text=SAMPLE_TEXT, user_id=dedup_user, db=db)
            # Second submission (same content)
            result2 = await memory_pipeline.run(text=SAMPLE_TEXT, user_id=dedup_user, db=db)
            elapsed = time.perf_counter() - t0
            dedup_works = result2.discarded_count > 0 or result2.new_count < result1.new_count
            details = (f"Round1: new={result1.new_count}, Round2: new={result2.new_count}, "
                       f"discarded={result2.discarded_count}")
            record("Phase4-DB", "dedup_with_db", "pass" if dedup_works else "fail",
                   elapsed * 1000, details=details)
            print(f"  {'✅' if dedup_works else '❌'} dedup_with_db ({elapsed:.1f}s): {details}")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "dedup_with_db", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ dedup_with_db: {e}")

    # ---- 4.3 MemoryStore Search ----
    print("\n--- 4.3 MemoryStore 语义搜索 ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_store import memory_store
        async for db in get_db():
            result = await memory_store.search(
                query="电商平台技术栈",
                user_id=TEST_USER_ID,
                db=db,
                top_k=5,
            )
            elapsed = time.perf_counter() - t0
            n_results = len(result.get("results", []))
            details = f"Found {n_results} results, total_candidates={result.get('total_candidates', 0)}, elapsed={elapsed:.1f}s"
            record("Phase4-DB", "memory_store_search", "pass", elapsed * 1000, details=details)
            print(f"  ✅ memory_store_search ({elapsed:.1f}s): {n_results} results")
            for r in result.get("results", [])[:3]:
                print(f"     [{r.get('memory_type', '?')}] score={r.get('relevance_score', 'N/A')} | {r.get('content', '')[:60]}")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "memory_store_search", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ memory_store_search: {e}")

    # ---- 4.4 MemoryStore CRUD ----
    print("\n--- 4.4 MemoryStore List & Context ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_store import memory_store
        async for db in get_db():
            # List
            list_result = await memory_store.list_memories(
                user_id=TEST_USER_ID, db=db, page=1, page_size=10,
            )
            n_list = list_result.get("total", 0)

            # Context
            ctx_result = await memory_store.get_context(
                query="项目进度和技术栈",
                user_id=TEST_USER_ID,
                db=db,
                max_tokens=1000,
            )
            n_ctx = ctx_result.get("memory_count", 0)
            formatted_len = len(ctx_result.get("formatted_text", ""))

            elapsed = time.perf_counter() - t0
            details = f"List: {n_list} total, Context: {n_ctx} memories, {formatted_len} chars formatted"
            record("Phase4-DB", "memory_store_crud", "pass", elapsed * 1000, details=details)
            print(f"  ✅ memory_store_crud ({elapsed:.1f}s): {details}")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "memory_store_crud", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ memory_store_crud: {e}")

    # ---- 4.5 MemoryStore Context Formatting ----
    print("\n--- 4.5 Context 格式化验证 ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_store import memory_store
        async for db in get_db():
            ctx = await memory_store.get_context(
                query="技术选型",
                user_id=TEST_USER_ID,
                db=db,
                max_tokens=2000,
                group_by_type=True,
            )
            formatted = ctx.get("formatted_text", "")
            elapsed = time.perf_counter() - t0
            has_grouping = "##" in formatted
            details = (f"Tokens: ~{ctx.get('estimated_tokens', 0)}, "
                       f"Grouped: {has_grouping}, Memories: {ctx.get('memory_count', 0)}")
            record("Phase4-DB", "memory_store_context", "pass", elapsed * 1000, details=details)
            print(f"  ✅ memory_store_context ({elapsed:.1f}s): {details}")
            # Show a preview
            preview = formatted[:300].replace("\n", "\n     ")
            print(f"     Preview:\n     {preview}...")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "memory_store_context", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ memory_store_context: {e}")

    # ---- 4.6 MemoryStore Update & Soft Delete ----
    print("\n--- 4.6 Update & Soft Delete ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_store import memory_store
        async for db in get_db():
            # First, get a memory ID from the pipeline result
            list_r = await memory_store.list_memories(user_id=TEST_USER_ID, db=db, page=1, page_size=1)
            items = list_r.get("items", [])
            if items:
                mem_id = items[0]["memory_id"]

                # Update
                update_r = await memory_store.update_memory(
                    memory_id=mem_id, db=db,
                    content="[UPDATED] " + items[0].get("content", ""),
                    importance=0.95,
                )
                assert update_r.get("updated"), f"Update failed: {update_r}"

                # Soft delete
                delete_r = await memory_store.soft_delete(
                    memory_id=mem_id, db=db, reason="测试清理",
                )
                assert delete_r.get("deleted"), f"Delete failed: {delete_r}"

                elapsed = time.perf_counter() - t0
                record("Phase4-DB", "memory_store_update_delete", "pass", elapsed * 1000,
                       details=f"Updated & deleted mem_id={mem_id[:12]}...")
                print(f"  ✅ memory_store_update_delete ({elapsed:.1f}s): mem_id={mem_id[:12]}...")
            else:
                record("Phase4-DB", "memory_store_update_delete", "skip",
                       details="No memories available to update/delete")
                print("  ⏭️ memory_store_update_delete: no memories available")
            break
    except Exception as e:
        elapsed = time.perf_counter() - t0
        record("Phase4-DB", "memory_store_update_delete", "fail", elapsed * 1000, error=str(e))
        print(f"  ❌ memory_store_update_delete: {e}")


# ====================================================================
# Phase 5: 降级与容错
# ====================================================================

async def phase5_degradation():
    """测试降级与容错场景。"""
    divider("Phase 5: 降级与容错场景")

    from app.services.memory_dedup import DedupService, DedupAction
    from app.services.embedding_client import EmbeddingClient
    from app.models.base import Memory

    mock_embedding = MagicMock(spec=EmbeddingClient)
    mock_embedding.embed_single = AsyncMock(return_value=[0.1] * 1024)

    # ---- 5.1 Qdrant Unavailable → All KEEP_NEW ----
    print("\n--- 5.1 Qdrant 不可用 → 全部 KEEP_NEW ---")
    t0 = time.perf_counter()
    try:
        mock_qdrant = MagicMock()
        mock_qdrant.is_available = False
        mock_qdrant.collection_name = "agent_mem_generation"
        service = DedupService(mock_embedding, mock_qdrant)
        db = AsyncMock()
        candidates = [
            MemoryCandidate(content="test1", summary="s1"),
            MemoryCandidate(content="test2", summary="s2"),
            MemoryCandidate(content="test3", summary="s3"),
        ]
        results = await service.process_candidates(candidates, "user_test", db)
        all_keep = all(r.action == DedupAction.KEEP_NEW for r in results)
        assert all_keep and len(results) == 3
        record("Phase5-Deg", "qdrant_unavailable_all_keep_new", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ qdrant_unavailable_all_keep_new: 3/3 KEEP_NEW")
    except Exception as e:
        record("Phase5-Deg", "qdrant_unavailable_all_keep_new", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ qdrant_unavailable_all_keep_new: {e}")

    # ---- 5.2 Empty Extraction → Empty Result ----
    print("\n--- 5.2 空抽取结果处理 ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_extractor import ExtractionResult
        from app.services.memory_generator import MemoryGenerator

        extraction = ExtractionResult(source_text="hello")
        mock_llm = MagicMock()
        generator = MemoryGenerator(mock_llm)
        candidates = await generator.generate(extraction)
        assert len(candidates) == 0
        record("Phase5-Deg", "empty_extraction_empty_result", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ empty_extraction_empty_result")
    except Exception as e:
        record("Phase5-Deg", "empty_extraction_empty_result", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ empty_extraction_empty_result: {e}")

    # ---- 5.3 MemoryCandidate Validation Edge Cases ----
    print("\n--- 5.3 MemoryCandidate 验证边界 ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_generator import MemoryCandidate

        # All valid types
        valid_types = ["fact", "preference", "task_state", "decision", "constraint", "process"]
        for vt in valid_types:
            c = MemoryCandidate(content="test", summary="s", memory_type=vt)
            c.validate()
            assert c.memory_type == vt

        # Edge values
        c = MemoryCandidate(content="test", summary="s", importance=0.0, confidence=1.0)
        c.validate()
        assert c.importance == 0.0
        assert c.confidence == 1.0

        # Negative
        c2 = MemoryCandidate(content="test", summary="s", importance=-100, confidence=999)
        c2.validate()
        assert c2.importance == 0.0
        assert c2.confidence == 1.0

        # Empty tags/entities
        c3 = MemoryCandidate(content="test", summary="s", key_points=[], tags=[], entities=[])
        c3.validate()
        assert c3.key_points == []
        assert c3.tags == []

        record("Phase5-Deg", "candidate_validation_edge_cases", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ candidate_validation_edge_cases")
    except Exception as e:
        record("Phase5-Deg", "candidate_validation_edge_cases", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ candidate_validation_edge_cases: {e}")

    # ---- 5.4 Keyword Extraction Edge Cases ----
    print("\n--- 5.4 关键词提取边界 ---")
    t0 = time.perf_counter()
    try:
        # English only
        nouns_en = DedupService._extract_nouns("Hello World Test")
        assert len(nouns_en) >= 2, f"Expected >=2 English keywords, got {nouns_en}"

        # Chinese only
        nouns_cn = DedupService._extract_nouns("你好世界测试用户偏好")
        assert len(nouns_cn) >= 1, f"Expected >=1 Chinese keywords, got {nouns_cn}"

        # Mixed
        nouns_mix = DedupService._extract_nouns("用户使用 Python 开发 Web 应用")
        assert len(nouns_mix) >= 2, f"Expected >=2 mixed keywords, got {nouns_mix}"

        # Empty
        nouns_empty = DedupService._extract_nouns("")
        assert nouns_empty == []

        record("Phase5-Deg", "keyword_extraction_edge_cases", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ keyword_extraction_edge_cases")
    except Exception as e:
        record("Phase5-Deg", "keyword_extraction_edge_cases", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ keyword_extraction_edge_cases: {e}")

    # ---- 5.5 DB-only Search Fallback ----
    print("\n--- 5.5 DB-only 搜索降级 ---")
    t0 = time.perf_counter()
    try:
        from app.services.memory_store import MemoryStore

        store = MemoryStore()
        mock_db = AsyncMock()

        # Simulate empty DB
        mock_result = MagicMock()
        mock_result.scalars = MagicMock()
        mock_result.scalars.return_value.all = MagicMock(return_value=[])
        mock_db.execute = AsyncMock(return_value=mock_result)

        result = await store._db_only_search(
            query="测试查询",
            user_id="test_user",
            db=mock_db,
            top_k=5,
        )
        assert result["fallback"] is True
        assert result["results"] == []
        assert result["total_candidates"] == 0

        record("Phase5-Deg", "db_only_search_fallback", "pass", (time.perf_counter() - t0) * 1000)
        print("  ✅ db_only_search_fallback")
    except Exception as e:
        record("Phase5-Deg", "db_only_search_fallback", "fail", (time.perf_counter() - t0) * 1000, error=str(e))
        print(f"  ❌ db_only_search_fallback: {e}")


# ====================================================================
# Word 报告生成
# ====================================================================

def generate_word_report(results: list[dict], output_path: str):
    """根据测试结果生成 Word 报告。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("  ⚠️ python-docx 未安装，跳过 Word 报告生成")
        print("     安装: pip install python-docx")
        return None

    def set_cell_shading(cell, color):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)

    def add_header_row(table, cells, color="1A5276"):
        row = table.rows[0]
        for i, text in enumerate(cells):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, color)

    def add_data_row(table, cells, row_idx):
        row = table.rows[row_idx]
        for i, text in enumerate(cells):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(text))
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F4F6F6")

    def add_code_block(doc, text):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return p

    # ---- Compute stats ----
    total = len(results)
    passed = sum(1 for r in results if r["status"] == "pass")
    failed = sum(1 for r in results if r["status"] == "fail")
    skipped = sum(1 for r in results if r["status"] == "skip")
    phases = {}
    for r in results:
        phase = r["phase"]
        if phase not in phases:
            phases[phase] = {"total": 0, "pass": 0, "fail": 0, "skip": 0}
        phases[phase]["total"] += 1
        phases[phase][r["status"]] += 1

    # ---- Build document ----
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # === Cover ===
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("记忆生成与去重融合模块\n全功能测试报告")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Memory Generation & Dedup — All-in-One Test Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = (f"项目: Agent Memory System (agent_mem)\n"
                 f"测试时间: {START_TIME.strftime('%Y-%m-%d %H:%M:%S')}\n"
                 f"总测试项: {total}  |  通过: {passed}  |  失败: {failed}  |  跳过: {skipped}\n"
                 f"通过率: {passed / max(total, 1) * 100:.1f}%")
    run = info.add_run(info_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x56, 0x6D, 0x7E)

    doc.add_page_break()

    # === Summary ===
    doc.add_heading("测试执行摘要", level=1)

    summary_table = doc.add_table(rows=len(phases) + 2, cols=5, style='Table Grid')
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(summary_table, ["阶段", "总数", "通过", "失败", "跳过"])

    phase_order = ["Phase1-LLM", "Phase1-Emb", "Phase1-Ext", "Phase1-Gen", "Phase1-Dedup",
                    "Phase2-Schema", "Phase2-Map",
                    "Phase3-E2E", "Phase4-DB", "Phase5-Deg"]
    phase_labels = {
        "Phase1-LLM": "P1 LLM Client", "Phase1-Emb": "P1 Embedding",
        "Phase1-Ext": "P1 MemoryExtractor", "Phase1-Gen": "P1 MemoryGenerator",
        "Phase1-Dedup": "P1 DedupService",
        "Phase2-Schema": "P2 Schema验证", "Phase2-Map": "P2 Pipeline映射",
        "Phase3-E2E": "P3 E2E基础", "Phase4-DB": "P4 E2E完整",
        "Phase5-Deg": "P5 降级容错",
    }

    row_idx = 1
    for p in phase_order:
        if p in phases:
            s = phases[p]
            add_data_row(summary_table, [
                phase_labels.get(p, p), s["total"], s["pass"], s["fail"], s["skip"]
            ], row_idx)
            row_idx += 1

    # Add total row
    add_data_row(summary_table, ["合计", total, passed, failed, skipped], row_idx)

    doc.add_paragraph()

    # === Detailed Results ===
    doc.add_heading("详细测试结果", level=1)

    for phase_name in phase_order:
        phase_results = [r for r in results if r["phase"] == phase_name]
        if not phase_results:
            continue

        label = phase_labels.get(phase_name, phase_name)
        doc.add_heading(label, level=2)

        n = len(phase_results)
        tbl = doc.add_table(rows=n + 1, cols=5, style='Table Grid')
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_header_row(tbl, ["#", "测试项", "状态", "耗时(ms)", "详情"])

        for i, r in enumerate(phase_results):
            status_icon = {"pass": "✅ PASS", "fail": "❌ FAIL", "skip": "⏭️ SKIP"}.get(r["status"], r["status"])
            details_text = r.get("details", "") or r.get("error", "") or ""
            if len(details_text) > 120:
                details_text = details_text[:117] + "..."
            add_data_row(tbl, [
                str(i + 1),
                r["name"],
                status_icon,
                str(r.get("duration_ms", 0)),
                details_text,
            ], i + 1)

        doc.add_paragraph()

    # === Known Issues & Notes ===
    doc.add_page_break()
    doc.add_heading("测试覆盖范围", level=1)

    doc.add_paragraph(
        "本测试脚本覆盖记忆生成与去重融合模块的以下核心功能："
    )

    coverage_items = [
        "LLM Client: 对话补全、结构化JSON抽取、Regex恢复、错误处理 — 4 项",
        "Embedding Client: 单文本嵌入(1024维)、批量拆分(32条/批次) — 2 项",
        "MemoryExtractor: 关键事实/任务状态/历史决策三路并行抽取 + Facade调度 + 序列化 — 5 项",
        "MemoryGenerator: 抽取→记忆候选转化、空抽取处理、候选验证与裁剪 — 3 项",
        "DedupService: 四类决策动作、关键词Jaccard、标识一致性检查、内容融合、Qdrant降级、中英文关键词提取 — 10 项",
        "Schema验证: 必填字段、role枚举、content非空、messages非空、user_id标准化、extraction_type校验 — 7 项",
        "Pipeline→Results映射: keep_new→ADD, merge→MERGE, discard→SKIP, update_existing→ADD — 1 项",
        "E2E基础: LLM/Embedding连接、真实抽取、真实生成、决策矩阵、边界评分 — 6 项",
        "E2E完整: 完整流水线、DB去重、语义搜索、CRUD操作、Context格式化、Update/SoftDelete — 6 项",
        "降级容错: Qdrant不可用、空抽取、候选验证边界、关键词提取边界、DB-only搜索 — 5 项",
    ]
    for item in coverage_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("测试环境", level=2)
    env_table = doc.add_table(rows=6, cols=2, style='Table Grid')
    add_header_row(env_table, ["项目", "值"])
    env_data = [
        ["Python", "3.12.13"],
        ["LLM", "DeepSeek (deepseek-chat)"],
        ["Embedding", "SiliconFlow BAAI/bge-m3 (1024维)"],
        ["数据库", "PostgreSQL (Docker: mem-postgres)"],
        ["向量库", "Qdrant (Docker: mem-qdrant, collection: agent_mem_generation)"],
    ]
    for i, (k, v) in enumerate(env_data):
        add_data_row(env_table, [k, v], i + 1)

    doc.add_paragraph()

    doc.add_heading("已知注意事项", level=2)
    notes = [
        "Phase 1-2 无需任何外部依赖，可在任何环境运行。适合 CI/CD 集成。",
        "Phase 3 需要 DeepSeek API + SiliconFlow API 连通。延迟约 5-15 秒。",
        "Phase 4 需要 Docker 容器 (mem-postgres + mem-qdrant) 运行中。",
        "Phase 5 为纯 Mock 测试，验证降级逻辑的正确性。",
        "使用 --mock-only 仅运行 Phase 1-2；使用 --no-db 仅运行 Phase 1-3。",
        "E2E 完整流水线每次 LLM 调用耗时 5-15 秒（4 次 LLM 调用：3 次抽取 + 1 次生成）。",
    ]
    for n in notes:
        doc.add_paragraph(n, style='List Bullet')

    # === Footer ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 文档结束 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
                     f"Agent Memory System v1.0.0  |  test_all_in_one.py")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

    doc.save(output_path)
    return output_path


# ====================================================================
# Main
# ====================================================================

def print_summary():
    """打印测试摘要。"""
    total = len(RESULTS)
    passed = sum(1 for r in RESULTS if r["status"] == "pass")
    failed = sum(1 for r in RESULTS if r["status"] == "fail")
    skipped = sum(1 for r in RESULTS if r["status"] == "skip")

    divider("测试摘要")
    print(f"  总测试项: {total}")
    print(f"  ✅ 通过:   {passed} ({passed / max(total, 1) * 100:.1f}%)")
    print(f"  ❌ 失败:   {failed}")
    print(f"  ⏭️ 跳过:   {skipped}")
    print(f"  总耗时:    {(datetime.now() - START_TIME).total_seconds():.1f}s")

    if failed > 0:
        print(f"\n  失败项:")
        for r in RESULTS:
            if r["status"] == "fail":
                print(f"    ❌ [{r['phase']}] {r['name']}: {r.get('error', r.get('details', ''))[:120]}")


async def main():
    print("=" * 70)
    print("  记忆生成与去重融合模块 — 一键全功能测试")
    print("=" * 70)
    print(f"  开始时间: {START_TIME.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"  测试用户: {TEST_USER_ID}")

    # Parse args
    args = set(sys.argv[1:])
    mock_only = "--mock-only" in args
    no_db = "--no-db" in args
    gen_report = "--report" in args

    if mock_only:
        print(f"  模式: Mock Only (Phase 1-2)")
    elif no_db:
        print(f"  模式: No-DB (Phase 1-3)")
    else:
        print(f"  模式: Full (Phase 1-5)")

    # Phase 1: Unit tests (always run)
    await phase1_unit_tests()

    # Phase 2: Schema & mapping (always run)
    await phase2_schema_and_mapping()

    if mock_only:
        print_summary()
        _save_and_report(gen_report)
        return

    # Phase 3: E2E basic
    e2e_ok = await phase3_e2e_basic()

    if no_db:
        print_summary()
        _save_and_report(gen_report)
        return

    # Phase 4: E2E full (needs DB + Qdrant)
    if e2e_ok:
        await phase4_e2e_full()
    else:
        print("\n  ⚠️ Phase 3 LLM 连接失败，跳过 Phase 4 (E2E Full)")

    # Phase 5: Degradation (always runs, no external deps needed)
    await phase5_degradation()

    # Summary
    print_summary()

    # Save & Report
    _save_and_report(gen_report)


def _save_and_report(gen_report: bool):
    """保存 JSON 结果并可选生成 Word 报告。"""
    # Save JSON
    json_path = Path(__file__).parent / "test_all_in_one_results.json"
    output_data = {
        "meta": {
            "start_time": START_TIME.isoformat(),
            "end_time": datetime.now().isoformat(),
            "total": len(RESULTS),
            "passed": sum(1 for r in RESULTS if r["status"] == "pass"),
            "failed": sum(1 for r in RESULTS if r["status"] == "fail"),
            "skipped": sum(1 for r in RESULTS if r["status"] == "skip"),
        },
        "results": RESULTS,
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    print(f"\n  📄 JSON 结果已保存: {json_path}")

    # Generate Word report
    if gen_report:
        print(f"\n  生成 Word 报告...")
        report_path = Path(__file__).parent / "记忆生成与去重融合_全功能测试报告.docx"
        result_path = generate_word_report(RESULTS, str(report_path))
        if result_path:
            print(f"  📄 Word 报告已生成: {result_path}")
    else:
        print(f"  💡 使用 --report 参数可同时生成 Word 报告")


if __name__ == "__main__":
    asyncio.run(main())
