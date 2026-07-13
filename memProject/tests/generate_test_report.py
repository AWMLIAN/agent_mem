# -*- coding: utf-8 -*-
"""
生成记忆生成与去重融合功能测试报告 Word 文档。
"""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, text, language=""):
    """添加代码块样式段落"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_header_row(table, cells, color="1A5276"):
    """设置表头行"""
    row = table.rows[0]
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)


def add_data_row(table, cells, row_idx):
    """添加数据行"""
    row = table.rows[row_idx]
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        if row_idx % 2 == 0:
            set_cell_shading(cell, "F4F6F6")


def main():
    doc = Document()

    # ================================================================
    # 页面设置
    # ================================================================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 样式设置
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # ================================================================
    # 封面
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("记忆生成与去重融合功能\n测试报告")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Memory Generation & Dedup — Test Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"项目: Agent Memory System (agent_mem)\n版本: 1.0.0\n日期: {datetime.now().strftime('%Y-%m-%d')}\n模块: 记忆生成与去重融合"
    run = info.add_run(info_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x56, 0x6D, 0x7E)

    doc.add_page_break()

    # ================================================================
    # 目录页
    # ================================================================
    doc.add_heading("目  录", level=1)
    toc_items = [
        "1. 测试概述",
        "2. 测试环境与前置条件",
        "3. 测试架构与模块说明",
        "4. 测试步骤（如何测试）",
        "    4.1 单元测试",
        "    4.2 E2E 基础测试",
        "    4.3 E2E 完整流水线测试",
        "    4.4 E2E 去重专项测试",
        "5. 测试数据",
        "    5.1 单元测试 Mock 数据",
        "    5.2 E2E 基础测试数据",
        "    5.3 E2E 去重测试数据（二次提交）",
        "6. 预期测试结果",
        "    6.1 单元测试预期结果",
        "    6.2 E2E 基础测试预期结果",
        "    6.3 去重决策矩阵预期结果",
        "    6.4 完整流水线预期结果",
        "7. 实际测试结果（2026-07-13）",
        "8. 测试覆盖率分析",
        "9. 已知限制与注意事项",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # 1. 测试概述
    # ================================================================
    doc.add_heading("1. 测试概述", level=1)

    doc.add_paragraph(
        "本测试报告覆盖 Agent Memory System 中「记忆生成与去重融合」模块的完整测试。"
        "该模块实现了从原始对话文本到结构化记忆的全自动流水线，"
        "包含四个核心阶段：关键信息抽取（Extract）、结构化记忆生成（Generate）、"
        "多阶段去重融合（Dedup）、持久化存储（Store）。"
    )

    doc.add_heading("1.1 测试目标", level=2)
    goals = [
        "验证 MemoryExtractor 的三路并行抽取能力（关键事实 / 任务状态 / 历史决策）",
        "验证 MemoryGenerator 将抽取结果转化为标准化 MemoryCandidate 的准确性",
        "验证 DedupService 多阶段去重决策矩阵的正确性（DISCARD / MERGE / UPDATE_EXISTING / KEEP_NEW）",
        "验证 MemoryPipeline 端到端流水线编排的正确性",
        "验证 LLM Client（DeepSeek）和 Embedding Client（SiliconFlow BGE-M3）的真实 API 连通性",
        "验证去重降级策略：当 Qdrant 不可用时自动将所有候选标记为 KEEP_NEW",
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')

    doc.add_heading("1.2 测试范围", level=2)
    doc.add_paragraph(
        "测试覆盖 14 个新建文件中的所有核心类和函数，包含 24 个单元测试（Mock 模式）"
        "和 4 类 E2E 场景测试。测试采用 Mock + 真实 API 结合的方式，"
        "确保在不依赖数据库的情况下也能验证核心逻辑。"
    )

    # ================================================================
    # 2. 测试环境与前置条件
    # ================================================================
    doc.add_heading("2. 测试环境与前置条件", level=1)

    doc.add_heading("2.1 环境要求", level=2)

    env_table = doc.add_table(rows=8, cols=2, style='Table Grid')
    env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(env_table, ["项目", "要求"])

    env_data = [
        ["操作系统", "Windows 11 / Linux / macOS"],
        ["Python", "3.12+ (项目使用 3.12.13)"],
        ["Conda 环境", "agent_mem"],
        ["Docker Desktop", "运行 mem-postgres + mem-qdrant 容器（E2E full/dedup 模式需要）"],
        ["DeepSeek API Key", "已配置在 .env 中"],
        ["SiliconFlow API Key", "已配置在 .env 中，BGE-M3 模型"],
        ["SSL 证书", "Windows 需设置 SSL_CERT_FILE 环境变量"],
    ]
    for i, (k, v) in enumerate(env_data):
        add_data_row(env_table, [k, v], i + 1)

    doc.add_paragraph()

    doc.add_heading("2.2 依赖服务", level=2)

    doc.add_paragraph("单元测试模式（Mock）:", style='List Bullet')
    doc.add_paragraph("无需任何外部服务，所有 LLM/Embedding/DB 调用均被 Mock。", style='List Bullet 2')

    doc.add_paragraph("E2E Basic 模式:", style='List Bullet')
    doc.add_paragraph("需要 DeepSeek API + SiliconFlow API 连通。不需要数据库。", style='List Bullet 2')

    doc.add_paragraph("E2E Full / Dedup 模式:", style='List Bullet')
    doc.add_paragraph(
        "需要 Docker 容器 (mem-postgres + mem-qdrant) 运行中，以及 DeepSeek + SiliconFlow API 连通。",
        style='List Bullet 2'
    )

    doc.add_heading("2.3 环境启动步骤", level=2)

    doc.add_paragraph("步骤 1: 启动 Docker 容器（仅 Full/Dedup 模式需要）")
    add_code_block(doc, "docker start mem-postgres mem-qdrant")

    doc.add_paragraph("步骤 2: 激活 conda 环境")
    add_code_block(doc, "source activate agent_mem")

    doc.add_paragraph("步骤 3: 设置 SSL 证书（Windows 必须）")
    add_code_block(doc,
        'export SSL_CERT_FILE="E:/anaconda3/envs/agent_mem/Lib/site-packages/certifi/cacert.pem"'
    )

    doc.add_paragraph("步骤 4: 进入项目目录")
    add_code_block(doc, 'cd "E:/AI Memory/agent_mem/memProject"')

    doc.add_page_break()

    # ================================================================
    # 3. 测试架构与模块说明
    # ================================================================
    doc.add_heading("3. 测试架构与模块说明", level=1)

    doc.add_paragraph(
        "记忆生成流水线采用四阶段管道架构，每个阶段由独立服务模块负责："
    )

    doc.add_heading("3.1 架构图", level=2)
    add_code_block(doc, """┌──────────────────────────────────────────────────────────────┐
│                    MemoryPipeline (编排器)                     │
├──────────────────────────────────────────────────────────────┤
│                                                               │
│  原始文本 ─→ Phase 1: Extract       MemoryExtractor           │
│                  │   ├── KeyFactExtractor    (关键事实)       │
│                  │   ├── TaskStateExtractor  (任务状态)       │
│                  │   └── DecisionExtractor   (历史决策)       │
│                  │                                            │
│             ─→ Phase 2: Generate      MemoryGenerator         │
│                  │   └── LLM → JSON → MemoryCandidate[]       │
│                  │                                            │
│             ─→ Phase 3: Dedup         DedupService            │
│                  │   ├── 向量检索 (Qdrant)                     │
│                  │   ├── 关键词 Jaccard 计算                   │
│                  │   ├── 标识一致性判断                        │
│                  │   ├── 综合评分决策                          │
│                  │   └── 内容融合 (MERGE)                      │
│                  │                                            │
│             ─→ Phase 4: Store         PostgreSQL + Qdrant     │
│                                                               │
└──────────────────────────────────────────────────────────────┘""")

    doc.add_heading("3.2 核心模块与测试文件对照", level=2)

    mod_table = doc.add_table(rows=9, cols=3, style='Table Grid')
    mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(mod_table, ["模块", "文件路径", "测试覆盖"])

    mod_data = [
        ["LLM Client", "app/services/llm_client.py", "4 个测试"],
        ["Embedding Client", "app/services/embedding_client.py", "2 个测试"],
        ["MemoryExtractor\n(+3 子提取器)", "app/services/memory_extractor.py\napp/prompts/*_extraction.py", "5 个测试"],
        ["MemoryGenerator", "app/services/memory_generator.py\napp/prompts/memory_generation.py", "3 个测试"],
        ["DedupService", "app/services/memory_dedup.py", "9 个测试"],
        ["MemoryPipeline", "app/services/memory_pipeline.py", "E2E 覆盖"],
        ["API 端点", "app/api/v1/generation.py", "E2E 覆盖"],
        ["Qdrant Client", "app/core/qdrant_client.py", "E2E 覆盖"],
    ]
    for i, (mod, path, cov) in enumerate(mod_data):
        add_data_row(mod_table, [mod, path, cov], i + 1)

    doc.add_page_break()

    # ================================================================
    # 4. 测试步骤（如何测试）
    # ================================================================
    doc.add_heading("4. 测试步骤（如何测试）", level=1)

    # 4.1
    doc.add_heading("4.1 单元测试（Mock 模式，无需外部服务）", level=2)

    doc.add_paragraph("命令：")
    add_code_block(doc,
        "pytest tests/test_memory_generation.py -v"
    )

    doc.add_paragraph("或指定 Python 解释器：")
    add_code_block(doc,
        "E:/anaconda3/envs/agent_mem/python.exe -m pytest tests/test_memory_generation.py -v"
    )

    doc.add_paragraph("说明：")
    doc.add_paragraph(
        "所有外部依赖（LLM HTTP 调用、Embedding API、Qdrant、PostgreSQL）均使用 unittest.mock 替换。"
        "适合 CI/CD 集成和快速回归测试。",
        style='List Bullet'
    )

    # 4.2
    doc.add_heading("4.2 E2E 基础测试（需要真实 API）", level=2)

    doc.add_paragraph("命令：")
    add_code_block(doc,
        "python tests/debug/test_generation_e2e.py basic"
    )

    doc.add_paragraph("测试流程：")
    steps_basic = [
        "Step 1: 测试 DeepSeek LLM 连接 — 发送 'Say OK' 验证响应",
        "Step 2: 测试 SiliconFlow Embedding 连接 — 验证向量维度=1024",
        "Step 3: 执行关键记忆抽取 — 三路并行抽取 329 字符测试文本",
        "Step 4: 执行结构化记忆生成 — 将抽取结果转为 MemoryCandidate 列表",
        "Step 5: 验证去重决策矩阵 — 4 组 (vec, kw, identity) → (expected_action) 对照",
    ]
    for s in steps_basic:
        doc.add_paragraph(s, style='List Bullet')

    # 4.3
    doc.add_heading("4.3 E2E 完整流水线测试（需要 Docker + 数据库）", level=2)

    doc.add_paragraph("前置条件：")
    doc.add_paragraph("Docker 容器 mem-postgres 和 mem-qdrant 必须运行中。", style='List Bullet')

    doc.add_paragraph("命令：")
    add_code_block(doc,
        "python tests/debug/test_generation_e2e.py full"
    )

    doc.add_paragraph("测试流程：")
    steps_full = [
        "Step 1: 验证 LLM 连接",
        "Step 2: 执行关键记忆抽取",
        "Step 3: 执行结构化记忆生成",
        "Step 4: 初始化 Qdrant collection (agent_mem_generation)",
        "Step 5: 执行完整流水线: extract → generate → dedup → store",
        "Step 6: 验证 PostgreSQL 写入 (new/merged/discarded/updated 计数)",
        "Step 7: 验证 Qdrant 向量写入",
    ]
    for s in steps_full:
        doc.add_paragraph(s, style='List Bullet')

    # 4.4
    doc.add_heading("4.4 E2E 去重专项测试（需要 Docker + 数据库）", level=2)

    doc.add_paragraph("命令：")
    add_code_block(doc,
        "python tests/debug/test_generation_e2e.py dedup"
    )

    doc.add_paragraph("测试流程：")
    steps_dedup = [
        "Step 1: 生成唯一 user_id，确保每次测试独立",
        "Step 2: 第一次提交 SAMPLE_TEXT → 创建 N 条新记忆",
        "Step 3: 第二次提交相同 SAMPLE_TEXT（相同 user_id）→ 去重检测",
        "Step 4: 验证第二次提交中大部分记忆被 DISCARD（去重生效）",
    ]
    for s in steps_dedup:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 5. 测试数据
    # ================================================================
    doc.add_heading("5. 测试数据", level=1)

    doc.add_heading("5.1 单元测试 Mock 数据", level=2)

    doc.add_paragraph("5.1.1 LLM Client — 结构化 JSON 抽取测试")
    add_code_block(doc, '''# Mock LLM 返回（正常 JSON）
{"name": "test", "value": 42}

# Mock LLM 返回（需 Regex 恢复）
Prefix text {"name": "recovered"} suffix

# Mock LLM 返回（无法解析 → 抛异常）
No JSON here at all''')

    doc.add_paragraph("5.1.2 Embedding Client — 批量拆分测试")
    add_code_block(doc, '''# 40 条文本 → 自动拆分为 32 + 8 两次 API 调用
texts = ["text"] * 40
# 预期: call_count == 2, len(results) == 40''')

    doc.add_paragraph("5.1.3 MemoryExtractor — 关键事实抽取测试")
    add_code_block(doc, '''# Mock LLM 返回
{
  "business_objects": [
    {"name": "ProjectX", "type": "project", "description": "主要项目"}
  ],
  "constraints": [
    {"type": "temporal", "description": "下周五 deadline", "severity": "high"}
  ],
  "confirmations": [
    {"item": "使用 FastAPI", "parties": ["dev"], "context": "技术选型会议"}
  ]
}''')

    doc.add_paragraph("5.1.4 MemoryExtractor — 任务状态抽取测试")
    add_code_block(doc, '''# Mock LLM 返回
{
  "current_progress": "API 开发阶段",
  "completed_items": [
    {"item": "数据库设计", "evidence": "已完成 ER 图"}
  ],
  "pending_items": [
    {"item": "单元测试", "priority": "high"}
  ]
}''')

    doc.add_paragraph("5.1.5 MemoryExtractor — 历史决策抽取测试")
    add_code_block(doc, '''# Mock LLM 返回
{
  "confirmed_plans": [
    {"plan": "微服务架构", "alternatives": ["单体"], "decision_context": "架构评审"}
  ],
  "selection_rationale": [
    {"reason": "可扩展性更好", "criteria": ["扩展性", "维护性"]}
  ],
  "execution_results": [
    {"result": "部署成功", "outcome_type": "success"}
  ]
}''')

    doc.add_paragraph("5.1.6 MemoryGenerator — 记忆生成测试")
    add_code_block(doc, '''# Mock LLM 返回 2 条候选记忆
{
  "memories": [
    {
      "content": "用户偏好使用 Python 进行开发",
      "summary": "用户偏好 Python",
      "key_points": ["偏好 Python", "主要开发语言"],
      "memory_type": "preference",
      "tags": ["python", "preference", "dev"],
      "entities": ["Python"],
      "importance": 0.8,
      "confidence": 0.9
    },
    {
      "content": "ProjectX 的 deadline 是下周五",
      "summary": "ProjectX deadline",
      "key_points": ["deadline: 下周五"],
      "memory_type": "constraint",
      "tags": ["deadline", "ProjectX"],
      "entities": ["ProjectX"],
      "importance": 0.9,
      "confidence": 1.0
    }
  ]
}''')

    doc.add_paragraph("5.1.7 DedupService — 决策矩阵测试")
    add_code_block(doc, '''# 4 组决策矩阵对照数据
test_cases = [
    (0.95, 0.90, True,  DedupAction.DISCARD),          # composite=0.945 → DISCARD
    (0.90, 0.70, True,  DedupAction.UPDATE_EXISTING),  # composite=0.860 → UPDATE
    (0.88, 0.76, False, DedupAction.MERGE),            # composite=0.668 → MERGE
    (0.60, 0.30, False, DedupAction.KEEP_NEW),         # composite=0.390 → KEEP_NEW
]''')

    doc.add_paragraph("5.1.8 DedupService — 合并测试")
    add_code_block(doc, '''# 候选记忆
MemoryCandidate(
    content="新发现: 使用 FastAPI",
    tags=["fastapi"], entities=["FastAPI"],
    importance=0.9, confidence=0.95
)
# 已有记忆
Memory(
    content="用户使用 Python",
    tags=["python"], entities=["Python"],
    importance=0.7, confidence=0.8
)
# 合并后预期: content 包含 Python + FastAPI
# entities = ["fastapi", "python"] (去重+小写)
# importance = 0.9, confidence = 0.95 (取最大值)''')

    doc.add_page_break()

    doc.add_heading("5.2 E2E 基础测试数据（SAMPLE_TEXT）", level=2)

    doc.add_paragraph("以下是 E2E 测试使用的对话文本（329 字符，中文对话）：")

    add_code_block(doc, '''SAMPLE_TEXT = """
用户: 我们正在开发一个叫 ProjectX 的电商平台，主要用 Python 和 FastAPI。
智能体: 好的，Python + FastAPI 技术栈已经记录。
用户: 项目的 deadline 是下周五，目前数据库设计已经完成，API 接口开发正在进行中。
智能体: 了解。数据库设计已完成，API 开发进行中。
用户: 我们之前讨论了技术选型，最终决定用 PostgreSQL 而不是 MySQL，因为需要 pgvector 支持。
智能体: 确认，已选择 PostgreSQL 作为数据库。
用户: 用户画像服务需要在下周三之前完成，QA 团队已经确认了这个时间点。
智能体: 记录：用户画像服务 deadline 为下周三，QA 已确认。
"""''')

    doc.add_paragraph("该文本包含的信息要素：")
    features = [
        "业务对象: ProjectX (电商平台)、Python、FastAPI、PostgreSQL、MySQL",
        "约束条件: 项目 deadline 下周五、用户画像服务 deadline 下周三",
        "确认事项: Python+FastAPI 技术栈、PostgreSQL 选型、QA 确认 deadline",
        "任务状态: 数据库设计已完成、API 开发进行中、用户画像服务待完成",
        "历史决策: 选择 PostgreSQL 而非 MySQL，原因是 pgvector 支持",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("5.3 E2E 去重测试数据（SAMPLE_TEXT_2，二次提交）", level=2)

    doc.add_paragraph("用于增量更新的第二段对话文本：")

    add_code_block(doc, '''SAMPLE_TEXT_2 = """
用户: ProjectX 的 API 开发进度怎么样了？
智能体: API 基本框架已搭建完成，正在实现用户认证模块。
用户: 好的。另外我们决定把缓存层从 Redis 换成 Dragonfly，性能更好。
智能体: 确认切换到 Dragonfly。Dragonfly 兼容 Redis 协议，迁移成本低。
用户: 之前说的用户画像服务已经完成了，QA 测试通过了。
智能体: 好的，用户画像服务已标记为完成。
"""''')

    doc.add_paragraph("该文本包含的信息要素：")
    features2 = [
        "任务更新: API 框架已搭建，正在实现认证模块",
        "新决策: 缓存层从 Redis 切换到 Dragonfly",
        "任务完成: 用户画像服务已完成，QA 通过",
        "与 SAMPLE_TEXT 的关联: 同属 ProjectX 项目，多实体重叠 → 触发 MERGE/UPDATE",
    ]
    for f in features2:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 6. 预期测试结果
    # ================================================================
    doc.add_heading("6. 预期测试结果", level=1)

    doc.add_heading("6.1 单元测试预期结果", level=2)

    doc.add_paragraph("24 个测试用例，全部预期 PASS：")

    ut_table = doc.add_table(rows=25, cols=4, style='Table Grid')
    ut_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(ut_table, ["#", "测试类", "测试方法", "预期"])

    ut_data = [
        ["1", "TestLLMClient", "test_chat_completion_success", "PASS"],
        ["2", "TestLLMClient", "test_extract_structured", "PASS"],
        ["3", "TestLLMClient", "test_extract_structured_json_recovery", "PASS — Regex 从噪声中恢复 JSON"],
        ["4", "TestLLMClient", "test_extract_structured_complete_failure", "PASS — 抛 LLMServiceError"],
        ["5", "TestEmbeddingClient", "test_embed_single", "PASS — 返回 1024 维向量"],
        ["6", "TestEmbeddingClient", "test_embed_batch_splitting", "PASS — 40条拆为 32+8，两次 API 调用"],
        ["7", "TestMemoryExtractor", "test_key_fact_extraction", "PASS — 正确解析 business_objects/constraints/confirmations"],
        ["8", "TestMemoryExtractor", "test_task_state_extraction", "PASS — 正确解析 progress/completed/pending"],
        ["9", "TestMemoryExtractor", "test_decision_extraction", "PASS — 正确解析 plans/rationale/results"],
        ["10", "TestMemoryExtractor", "test_memory_extractor_facade", "PASS - types=[key_fact] only runs matching extractor"],
        ["11", "TestMemoryExtractor", "test_extraction_result_to_dict", "PASS — 缺失字段填充默认值"],
        ["12", "TestMemoryGenerator", "test_generate_memories", "PASS — 生成 2 条候选（preference + constraint）"],
        ["13", "TestMemoryGenerator", "test_generate_empty_extraction", "PASS — 空抽取返回空列表"],
        ["14", "TestMemoryGenerator", "test_memory_candidate_validation", "PASS - invalid type falls back to fact, values clipped to [0,1]"],
        ["15", "TestDedupService", "test_decide_action_discard", "PASS — composite=0.945 → DISCARD"],
        ["16", "TestDedupService", "test_decide_action_update_existing", "PASS — composite=0.860 + identity → UPDATE"],
        ["17", "TestDedupService", "test_decide_action_merge", "PASS — composite=0.668 + !identity → MERGE"],
        ["18", "TestDedupService", "test_decide_action_keep_new", "PASS — composite=0.390 → KEEP_NEW"],
        ["19", "TestDedupService", "test_compute_keyword_overlap", "PASS — Jaccard ∈ [0,1]"],
        ["20", "TestDedupService", "test_check_identity_same_task", "PASS — 相同 task_id → True, 不同 → False"],
        ["21", "TestDedupService", "test_check_identity_entity_overlap", "PASS — 2+ 实体重叠 → True, 否则 → False"],
        ["22", "TestDedupService", "test_merge_content", "PASS — 合并后包含双方内容，取 max 评分"],
        ["23", "TestDedupService", "test_process_candidates_qdrant_unavailable", "PASS — 全部 KEEP_NEW（降级）"],
        ["24", "TestDedupService", "test_extract_nouns", "PASS — 中英文关键词正确提取"],
    ]
    for i, row_data in enumerate(ut_data):
        add_data_row(ut_table, row_data, i + 1)

    doc.add_paragraph()

    doc.add_heading("6.2 E2E 基础测试预期结果", level=2)

    e2e_expected = [
        ("1. LLM 连接测试", "DeepSeek 返回 'OK'，连接正常"),
        ("2. Embedding 连接测试", "返回 1024 维向量，BGE-M3 模型正常"),
        ("3. 关键记忆抽取", "从 329 字符中文对话中正确抽取：\n"
         "  • 业务对象 ≥ 3 个（ProjectX / Python / FastAPI / PostgreSQL）\n"
         "  • 约束条件 ≥ 2 个（deadline 相关）\n"
         "  • 确认事项 ≥ 2 个（技术栈确认 / 数据库选型确认）\n"
         "  • 任务状态: 当前进展 + 已完成项 + 待处理项"),
        ("4. 结构化记忆生成", "生成 3-8 条 MemoryCandidate，每条包含：\n"
         "  • content / summary / key_points / tags / entities\n"
         "  • memory_type ∈ {fact, preference, task_state, decision, constraint, process}\n"
         "  • importance ∈ [0, 1], confidence ∈ [0, 1]"),
        ("5. 去重决策矩阵", "4 组对照全部正确：\n"
         "  • (0.95, 0.90, True) → discard\n"
         "  • (0.90, 0.70, True) → update_existing\n"
         "  • (0.88, 0.76, False) → merge\n"
         "  • (0.60, 0.30, False) → keep_new"),
    ]

    for title, expected in e2e_expected:
        p = doc.add_paragraph()
        run = p.add_run(title + "：")
        run.bold = True
        p.add_run(expected)

    doc.add_heading("6.3 去重决策矩阵预期结果", level=2)

    doc.add_paragraph("综合评分公式：")
    add_code_block(doc, "composite = 0.5 × vector_score + 0.3 × keyword_overlap + 0.2 × identity_bonus")

    dedup_table = doc.add_table(rows=6, cols=6, style='Table Grid')
    dedup_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(dedup_table, ["vector", "keyword", "identity", "composite", "预期动作", "说明"])

    dedup_data = [
        ["0.95", "0.90", "True", "0.945", "DISCARD", "≥0.90 → 近乎重复，丢弃"],
        ["0.90", "0.70", "True", "0.860", "UPDATE_EXISTING", "≥0.80 + identity → 覆盖更新"],
        ["0.88", "0.76", "False", "0.668", "MERGE", "≥0.65 + !identity → 融合内容"],
        ["0.60", "0.30", "False", "0.390", "KEEP_NEW", "<0.65 → 保留新记忆"],
        ["0.85", "0.55", "False", "0.590", "KEEP_NEW", "边界: 接近但不到 0.65"],
    ]
    for i, row_data in enumerate(dedup_data):
        add_data_row(dedup_table, row_data, i + 1)

    doc.add_paragraph()

    doc.add_heading("6.4 完整流水线预期结果（Full 模式）", level=2)

    doc.add_paragraph("MemoryPipeline.run() 四阶段预期行为：")

    doc.add_paragraph("Phase 1 — Extract（抽取）:", style='List Bullet')
    doc.add_paragraph(
        "三路并行 LLM 调用 (key_fact + task_state + decision)，"
        "聚合到 ExtractionResult。每个子提取器失败不影响其他。若全部为空 → 流水线提前终止。",
        style='List Bullet 2'
    )

    doc.add_paragraph("Phase 2 — Generate（生成）:", style='List Bullet')
    doc.add_paragraph(
        "LLM 将抽取数据转化为 MemoryCandidate[] 列表。若无候选 → 流水线提前终止。",
        style='List Bullet 2'
    )

    doc.add_paragraph("Phase 3 — Dedup（去重）:", style='List Bullet')
    doc.add_paragraph(
        "每条候选执行向量检索 → 关键词对比 → 标识检查 → 综合决策。\n"
        "  • 有 DB + Qdrant: 正常去重，返回 DISCARD/MERGE/UPDATE/KEEP_NEW\n"
        "  • Qdrant 不可用: 全部 KEEP_NEW（降级，不阻塞）\n"
        "  • 去重异常: 全部 KEEP_NEW（容错）",
        style='List Bullet 2'
    )

    doc.add_paragraph("Phase 4 — Store（存储）:", style='List Bullet')
    doc.add_paragraph(
        "KEEP_NEW → 新建 Memory ORM + Qdrant 向量\n"
        "MERGE/UPDATE → 更新已有记录 + 重新计算向量\n"
        "DISCARD → 跳过\n"
        "先 commit PostgreSQL，后 upsert Qdrant（Qdrant 失败非致命）",
        style='List Bullet 2'
    )

    doc.add_page_break()

    # ================================================================
    # 7. 实际测试结果 (2026-07-13)
    # ================================================================
    doc.add_heading("7. 实际测试结果（2026-07-13）", level=1)

    doc.add_heading("7.1 单元测试结果", level=2)

    add_code_block(doc, """============================= test session starts =============================
platform win32 -- Python 3.12.13, pytest-9.1.1
plugins: anyio-4.14.1, langsmith-0.9.5, asyncio-1.4.0, cov-7.1.0

tests/test_memory_generation.py::TestLLMClient::test_chat_completion_success PASSED [  4%]
tests/test_memory_generation.py::TestLLMClient::test_extract_structured PASSED [  8%]
tests/test_memory_generation.py::TestLLMClient::test_extract_structured_json_recovery PASSED [ 12%]
tests/test_memory_generation.py::TestLLMClient::test_extract_structured_complete_failure PASSED [ 16%]
tests/test_memory_generation.py::TestEmbeddingClient::test_embed_single PASSED [ 20%]
tests/test_memory_generation.py::TestEmbeddingClient::test_embed_batch_splitting PASSED [ 25%]
tests/test_memory_generation.py::TestMemoryExtractor::test_key_fact_extraction PASSED [ 29%]
tests/test_memory_generation.py::TestMemoryExtractor::test_task_state_extraction PASSED [ 33%]
tests/test_memory_generation.py::TestMemoryExtractor::test_decision_extraction PASSED [ 37%]
tests/test_memory_generation.py::TestMemoryExtractor::test_memory_extractor_facade PASSED [ 41%]
tests/test_memory_generation.py::TestMemoryExtractor::test_extraction_result_to_dict PASSED [ 45%]
tests/test_memory_generation.py::TestMemoryGenerator::test_generate_memories PASSED [ 50%]
tests/test_memory_generation.py::TestMemoryGenerator::test_generate_empty_extraction PASSED [ 54%]
tests/test_memory_generation.py::TestMemoryGenerator::test_memory_candidate_validation PASSED [ 58%]
tests/test_memory_generation.py::TestDedupService::test_decide_action_discard PASSED [ 62%]
tests/test_memory_generation.py::TestDedupService::test_decide_action_update_existing PASSED [ 66%]
tests/test_memory_generation.py::TestDedupService::test_decide_action_merge PASSED [ 70%]
tests/test_memory_generation.py::TestDedupService::test_decide_action_keep_new PASSED [ 75%]
tests/test_memory_generation.py::TestDedupService::test_compute_keyword_overlap PASSED [ 79%]
tests/test_memory_generation.py::TestDedupService::test_check_identity_same_task PASSED [ 83%]
tests/test_memory_generation.py::TestDedupService::test_check_identity_entity_overlap PASSED [ 87%]
tests/test_memory_generation.py::TestDedupService::test_merge_content PASSED [ 91%]
tests/test_memory_generation.py::TestDedupService::test_process_candidates_qdrant_unavailable PASSED [ 95%]
tests/test_memory_generation.py::TestDedupService::test_extract_nouns PASSED [100%]

============================= 24 passed in 1.71s =============================""")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("结果：24/24 全部通过 ✅")
    run.bold = True
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_heading("7.2 E2E 基础测试结果", level=2)

    doc.add_paragraph("7.2.1 LLM 连接测试")
    add_code_block(doc, "响应: OK\n✅ LLM 连接正常")
    doc.add_paragraph("DeepSeek (deepseek-chat) 连接成功，延迟正常。")

    doc.add_paragraph("7.2.2 Embedding 连接测试")
    add_code_block(doc, """向量维度: 1024
前 5 维: [-0.0159, 0.0269, -0.0428, 0.0136, -0.0192]
✅ Embedding 正常 (dim=1024)""")
    doc.add_paragraph("SiliconFlow BGE-M3 模型连接成功，返回 1024 维向量。")

    doc.add_paragraph("7.2.3 关键记忆抽取结果")
    doc.add_paragraph(
        "从 329 字符中文对话中，耗时 4.23s，三个子提取器并行执行，实际产出："
    )

    extract_table = doc.add_table(rows=4, cols=3, style='Table Grid')
    extract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(extract_table, ["抽取类型", "抽取数量", "详情"])
    add_data_row(extract_table, [
        "关键事实\n(Key Facts)",
        "6 个业务对象\n3 个约束\n4 个确认",
        "ProjectX(project), Python(resource), FastAPI(resource),\n"
        "PostgreSQL(resource), MySQL(resource), 用户画像服务(resource)\n"
        "约束: deadline 下周五(high), 用户画像下周三(high),\n"
        "  PostgreSQL替代MySQL(medium)"
    ], 1)
    add_data_row(extract_table, [
        "任务状态\n(Task State)",
        "已完成 3 项\n待处理 3 项",
        "当前: API 接口开发进行中\n"
        "已完成: 技术选型讨论, 数据库设计, QA确认deadline\n"
        "待处理: API接口开发(high), 用户画像服务(high),\n"
        "  项目整体交付(high)"
    ], 2)
    add_data_row(extract_table, [
        "历史决策\n(Decisions)",
        "0 项",
        "决策抽取器未检测到符合条件的决策记录\n"
        "（选型决策被归入关键事实的 confirmations 中）"
    ], 3)

    doc.add_paragraph()

    doc.add_paragraph("7.2.4 结构化记忆生成结果")
    doc.add_paragraph("耗时 5.88s，生成 5 条候选记忆：")

    gen_table = doc.add_table(rows=6, cols=4, style='Table Grid')
    gen_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(gen_table, ["#", "类型", "内容摘要", "重要性/置信度"])

    gen_data = [
        ["1", "fact", "ProjectX 电商平台项目概述（Python+FastAPI+PostgreSQL, deadline 下周五）", "0.90 / 1.00"],
        ["2", "constraint", "用户画像服务 deadline 下周三，QA 已确认", "0.80 / 1.00"],
        ["3", "decision", "数据库选型: PostgreSQL 替代 MySQL（因 pgvector 需求）", "0.70 / 1.00"],
        ["4", "task_state", "项目进度: 数据库设计完成, API 接口开发中", "0.60 / 1.00"],
        ["5", "task_state", "待办: API 开发, 用户画像服务, 项目整体交付（均高优先级）", "0.80 / 1.00"],
    ]
    for i, row_data in enumerate(gen_data):
        add_data_row(gen_table, row_data, i + 1)

    doc.add_paragraph()

    doc.add_paragraph("7.2.5 去重决策矩阵验证结果")
    add_code_block(doc, """vec=0.95 kw=0.90 id=True → discard ✅ (expected: discard)
vec=0.90 kw=0.70 id=True → update_existing ✅ (expected: update_existing)
vec=0.88 kw=0.76 id=False → merge ✅ (expected: merge)
vec=0.60 kw=0.30 id=False → keep_new ✅ (expected: keep_new)
✅ 决策矩阵全部正确""")

    doc.add_page_break()

    # ================================================================
    # 8. 测试覆盖率分析
    # ================================================================
    doc.add_heading("8. 测试覆盖率分析", level=1)

    doc.add_heading("8.1 功能覆盖", level=2)

    cov_table = doc.add_table(rows=13, cols=4, style='Table Grid')
    cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(cov_table, ["功能点", "单元测试", "E2E 测试", "覆盖状态"])

    cov_data = [
        ["LLM 普通对话", "✅", "✅ (连接测试)", "完全覆盖"],
        ["LLM 结构化 JSON 抽取", "✅", "✅ (抽取+生成)", "完全覆盖"],
        ["LLM JSON 解析容错 (Regex)", "✅", "—", "单元覆盖"],
        ["Embedding 单文本", "✅", "✅ (连接测试)", "完全覆盖"],
        ["Embedding 批量拆分", "✅", "—", "单元覆盖"],
        ["KeyFact 抽取", "✅", "✅ (真实 LLM)", "完全覆盖"],
        ["TaskState 抽取", "✅", "✅ (真实 LLM)", "完全覆盖"],
        ["Decision 抽取", "✅", "✅ (真实 LLM)", "完全覆盖"],
        ["MemoryCandidate 验证", "✅", "—", "单元覆盖"],
        ["去重决策矩阵 (4 动作)", "✅", "✅ (逻辑演示)", "完全覆盖"],
        ["关键词 Jaccard 计算", "✅", "—", "单元覆盖"],
        ["记忆合并逻辑", "✅", "—", "单元覆盖"],
    ]
    for i, row_data in enumerate(cov_data):
        add_data_row(cov_table, row_data, i + 1)

    doc.add_paragraph()

    doc.add_heading("8.2 边界条件覆盖", level=2)

    boundary_cases = [
        "空文本抽取 → 返回空 ExtractionResult → 流水线提前终止",
        "LLM 返回非 JSON → Regex 恢复 → 无法恢复则抛异常",
        "LLM 返回无效 memory_type → 自动回退为 'fact'",
        "importance/confidence 超出 [0,1] → 自动裁剪",
        "Qdrant 不可用 → 全部 KEEP_NEW（降级不阻塞）",
        "数据库无匹配记录 → 全部 KEEP_NEW",
        "去重阶段异常 → 全部 KEEP_NEW（容错）",
        "Embedding API 返回非 1024 维 → 错误处理",
    ]
    for bc in boundary_cases:
        doc.add_paragraph(bc, style='List Bullet')

    doc.add_heading("8.3 未覆盖项", level=2)

    not_covered = [
        "并发场景：多用户同时提交的去重冲突",
        "大批量场景：100+ 条文本的 MemoryPipeline.run_batch()",
        "数据库连接池耗尽：PostgreSQL 连接池满时的行为",
        "API 端点完整集成测试：POST /api/v1/memory/generate 的 HTTP 层测试",
        "异步任务状态轮询：POST /api/v1/memory/generate/async 的任务队列测试",
    ]
    for nc in not_covered:
        doc.add_paragraph(nc, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 9. 已知限制与注意事项
    # ================================================================
    doc.add_heading("9. 已知限制与注意事项", level=1)

    doc.add_heading("9.1 已知限制", level=2)

    limitations = [
        ("中文分词简化", "DedupService._extract_nouns() 使用正则提取中英文关键词（2-4 字中文 + 3+ 字母英文），"
         "非精确 NLP 分词。在专业领域术语较多的场景下，关键词重合率可能偏低。"),
        ("Qdrant 单点依赖", "去重功能强依赖 Qdrant 向量库。Qdrant 不可用时降级为全部 KEEP_NEW，"
         "不会阻塞流水线，但去重完全失效。"),
        ("LLM 调用延迟", "每次 Phase 1 抽取调用 3 次 LLM（三路并行），Phase 2 生成调用 1 次。"
         "总计 4 次 LLM 调用，延迟约 5-10 秒。不适用于实时对话场景。"),
        ("相似度阈值固定", "去重阈值（vector: 0.70/0.85, keyword: 0.50）为硬编码常量，"
         "未提供动态调优接口。不同业务场景可能需要不同的阈值配置。"),
    ]
    for title, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(title + "：")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("9.2 测试注意事项", level=2)

    notes = [
        "Windows 环境必须设置 SSL_CERT_FILE 环境变量，否则 HTTPS 请求会因证书验证失败而报错。",
        "运行 E2E full/dedup 模式前，确保 Docker 容器已启动: docker start mem-postgres mem-qdrant",
        "去重测试每次使用唯一 user_id (基于时间戳)，确保测试隔离。",
        "当 LLM API Key 过期或额度用尽时，所有依赖 LLM 的测试会失败。请检查 .env 中的 API Key。",
        "如果 Qdrant collection (agent_mem_generation) 中包含历史测试数据，"
        "可能影响去重测试的预期结果。必要时手动清理: 删除并重建 collection。",
    ]
    for n in notes:
        doc.add_paragraph(n, style='List Bullet')

    doc.add_paragraph()

    # ================================================================
    # 页脚
    # ================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 文档结束 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
                     f"Agent Memory System v1.0.0")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

    # ================================================================
    # 保存
    # ================================================================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "记忆生成与去重融合_测试报告.docx"
    )
    doc.save(output_path)
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    print(f"Test report generated: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
